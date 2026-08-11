from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "linkgen-public-operation-plan.md"
OUTPUT = ROOT / "docs" / "LinkGen-公开运营准备与执行路线-公司主体版.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT = "Calibri"
EAST_ASIA_FONT = "PingFang SC"
INK = "24323D"
MUTED = "687680"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
ORANGE = "F26430"
LIGHT_BLUE = "E8EEF5"
LIGHT_ORANGE = "FFF1EB"
BORDER = "C9D3DB"
def set_run_font(run, size=11, color=INK, bold=None, italic=None, name=FONT):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:eastAsia"), "zh-CN")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, color=ORANGE, size="16", space="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    borders.append(left)


def add_hyperlink(paragraph, label, url):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\([^\)]+\)|`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*)")


def add_inline(paragraph, text, size=11, color=INK):
    text = text.replace("<br>", "\n")
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        link = re.fullmatch(r"\[([^\]]+)\]\(([^\)]+)\)", token)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2))
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size - 0.5, color=DARK_BLUE, name="Courier New")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def split_table_row(line):
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_table_separator(line):
    cells = split_table_row(line)
    return len(cells) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def table_widths(column_count):
    options = {
        2: [2700, 6660],
        3: [1900, 3000, 4460],
        4: [1500, 2300, 2800, 2760],
        5: [1300, 1900, 1900, 2100, 2160],
    }
    if column_count in options:
        return options[column_count]
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    set_table_geometry(table, table_widths(len(rows[0])))
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            if row_index == 0:
                add_inline(paragraph, value, size=9.5, color=DARK_BLUE)
                for run in paragraph.runs:
                    run.bold = True
            else:
                add_inline(paragraph, value, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_border(paragraph)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_ORANGE)
    p_pr.append(shd)
    add_inline(paragraph, text, size=10.5, color=DARK_BLUE)


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.05
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    p_pr.append(shd)
    run = paragraph.add_run(code)
    set_run_font(run, size=9, color=DARK_BLUE, name="Courier New")


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    decimal_abstract_id = None
    abstract_ids = []
    for abstract in numbering.findall(qn("w:abstractNum")):
        abstract_id = int(abstract.get(qn("w:abstractNumId")))
        abstract_ids.append(abstract_id)
        level = abstract.find(qn("w:lvl"))
        if level is None or level.get(qn("w:ilvl")) != "0":
            continue
        num_fmt = level.find(qn("w:numFmt"))
        if num_fmt is not None and num_fmt.get(qn("w:val")) == "decimal":
            decimal_abstract_id = abstract_id
            break
    if decimal_abstract_id is None:
        decimal_abstract_id = max(abstract_ids or [0]) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(decimal_abstract_id))
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), "%1.")
        level.append(start)
        level.append(num_fmt)
        level.append(level_text)
        abstract.append(level)
        numbering.append(abstract)
    existing_num_ids = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    num_id = max(existing_num_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(decimal_abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_page_furniture(section):
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.paragraph_format.space_after = Pt(4)
    run = header_paragraph.add_run("LINKGEN  /  公开运营执行版")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(4)
    run = footer_paragraph.add_run("公司主体版  ·  2026-08-11  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer_paragraph._p.append(page_field)


def add_masthead(doc, title, metadata):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(7)
    run = kicker.add_run("LINKGEN  /  OPERATING PLAN")
    set_run_font(run, size=9, color=ORANGE, bold=True)

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(5)
    title_paragraph.paragraph_format.keep_with_next = True
    run = title_paragraph.add_run(title)
    set_run_font(run, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("公司主体申请、真实数据和公开运营的产品与工程执行路线")
    set_run_font(run, size=12, color=MUTED)

    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=MUTED)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def parse_document(doc, lines):
    title = lines[0].lstrip("# ").strip()
    metadata = []
    index = 1
    metadata_labels = ("调研日期", "适用产品", "主体方案", "范围")
    while index < len(lines):
        raw = lines[index].strip()
        if not raw:
            index += 1
            continue
        matched = False
        for label in metadata_labels:
            prefix = f"{label}："
            if raw.startswith(prefix):
                metadata.append((label, raw[len(prefix):].replace("<br>", "").strip()))
                matched = True
                break
        if matched:
            index += 1
            continue
        break

    add_masthead(doc, title, metadata)
    in_code = False
    code_lines = []
    previous_list_kind = None
    current_num_id = None

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            previous_list_kind = None
            index += 1
            continue
        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            previous_list_kind = None
            index += 1
            continue
        if stripped.startswith("## "):
            paragraph = doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            previous_list_kind = None
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:].strip())
            previous_list_kind = None
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows = [split_table_row(line), split_table_row(lines[index + 1])]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            rows.pop(1)
            add_table(doc, rows)
            previous_list_kind = None
            continue
        ordered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if ordered or bullet:
            style = "List Number" if ordered else "List Bullet"
            content = ordered.group(2) if ordered else bullet.group(1)
            if content.startswith("[ ] "):
                content = "☐ " + content[4:]
            elif content.startswith("[x] ") or content.startswith("[X] "):
                content = "☑ " + content[4:]
            paragraph = doc.add_paragraph(style=style)
            if ordered:
                if previous_list_kind != "ordered":
                    current_num_id = create_decimal_numbering(doc)
                apply_numbering(paragraph, current_num_id)
            add_inline(paragraph, content, size=11, color=INK)
            previous_list_kind = "ordered" if ordered else "bullet"
            index += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        add_inline(paragraph, stripped, size=11, color=INK)
        previous_list_kind = None
        index += 1


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    add_page_furniture(section)
    parse_document(doc, lines)
    doc.core_properties.title = "LinkGen 公开运营准备与执行路线（公司主体版）"
    doc.core_properties.subject = "LinkGen AI 社群微信小程序公开运营路线"
    doc.core_properties.author = "LinkGen"
    doc.core_properties.keywords = "LinkGen, 微信小程序, 公司主体, 公开运营, CloudBase"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
